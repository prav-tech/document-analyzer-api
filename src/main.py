import re
import base64
import os
import tempfile
from fastapi import FastAPI, Header, HTTPException
from pydantic import BaseModel
import fitz
from docx import Document
from PIL import Image, ImageFilter, ImageEnhance
import pytesseract
import numpy as np

# -------------------------------
# CONFIG
# -------------------------------
if os.name == "nt":  # Windows
    pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"

app = FastAPI()

# -------------------------------
# REQUEST MODEL
# -------------------------------
class DocumentRequest(BaseModel):
    fileName: str
    fileType: str
    fileBase64: str

# -------------------------------
# TEXT EXTRACTION
# -------------------------------
def extract_text(file_path, file_type):
    text = ""
    try:
        if file_type == "pdf":
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text()

        elif file_type == "docx":
            doc = Document(file_path)
            for para in doc.paragraphs:
                text += para.text + "\n"
            # Also extract tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                text += "\n"

        elif file_type == "image":
            try:
                img = Image.open(file_path)

                # Convert to RGB first if needed (handles PNG transparency etc.)
                if img.mode != 'RGB':
                    img = img.convert('RGB')

                # Convert to grayscale
                img = img.convert('L')

                # Resize if image is too small (improves OCR accuracy)
                width, height = img.size
                if width < 1000:
                    scale = 1000 / width
                    img = img.resize(
                        (int(width * scale), int(height * scale)),
                        Image.LANCZOS
                    )

                # Enhance contrast before thresholding
                enhancer = ImageEnhance.Contrast(img)
                img = enhancer.enhance(2.0)

                # Sharpen the image
                img = img.filter(ImageFilter.SHARPEN)

                # Convert to numpy array and apply threshold
                img_np = np.array(img)
                img_np = (img_np > 150) * 255
                img = Image.fromarray(img_np.astype('uint8'))

                # OCR with better config
                text = pytesseract.image_to_string(
                    img,
                    config='--oem 3 --psm 6'
                )

                # Clean text
                text = text.strip()

                if not text:
                    # Try again with different psm mode
                    text = pytesseract.image_to_string(
                        img,
                        config='--oem 3 --psm 3'
                    ).strip()

                if not text:
                    text = "No text found in image"

            except Exception as e:
                text = f"Error processing image: {str(e)}"

    except Exception as e:
        text = f"Error extracting text: {str(e)}"

    return text

# -------------------------------
# SENTIMENT
# -------------------------------
def get_sentiment(text):
    positive = {
        "good", "excellent", "profit", "success", "approved", "paid",
        "great", "outstanding", "achieved", "positive", "increase",
        "growth", "benefit", "award", "congratulations", "happy"
    }
    negative = {
        "loss", "failed", "rejected", "overdue", "penalty", "dispute",
        "unpaid", "error", "issue", "problem", "delay", "cancel",
        "terminate", "complaint", "violation", "damage"
    }

    words = set(text.lower().split())
    pos_count = len(words & positive)
    neg_count = len(words & negative)

    if pos_count > neg_count:
        return "Positive"
    elif neg_count > pos_count:
        return "Negative"
    else:
        return "Neutral"

# -------------------------------
# ENTITY EXTRACTION
# -------------------------------
def extract_entities(text):
    entities = {
        "names": [],
        "dates": [],
        "organizations": [],
        "amounts": [],
        "locations": [],
        "emails": [],
        "phones": [],
        "websites": [],
        "skills": [],
        "designations": []
    }

    clean_text = text.replace("\n", " ")

    blacklist = [
        "Design", "Designer", "Art", "Graphic", "Creative",
        "Media", "Campaign", "Suite", "Installations",
        "Agency", "Brand", "Web", "Photoshop", "Illustrator",
        "Profile", "Experience", "Education", "Portfolio",
        "Contact", "Skills", "Interests", "Present"
    ]

    known_locations = [
        "New York", "Visakhapatnam", "Hyderabad", "Bangalore",
        "Chennai", "Delhi", "Mumbai", "Andhra Pradesh", "Telangana",
        "Brooklyn", "NY", "Los Angeles", "San Francisco"
    ]

    known_skills = [
        "Photoshop", "Illustrator", "Web Design", "Figma",
        "Adobe Creative Suite", "Python", "Java", "JavaScript",
        "React", "Node.js", "SQL", "Machine Learning", "AutoCAD"
    ]

    known_designations = [
        "Graphic Designer", "Senior Graphic Designer", "Software Engineer",
        "Data Scientist", "Product Manager", "Web Developer",
        "UI/UX Designer", "Marketing Manager", "Business Analyst",
        "Project Manager", "Full Stack Developer", "DevOps Engineer"
    ]

    # NAMES
    name_pattern = re.findall(
        r'\b([A-Z][a-z]+\s[A-Z][a-z]+(?:\s[A-Z][a-z]+)?)\b',
        clean_text
    )
    filtered_names = []
    for name in name_pattern:
        words = name.split()
        if (
            not any(word in name for word in blacklist)
            and name not in known_locations
            and name not in known_skills
            and name not in known_designations
            and 2 <= len(words) <= 3
            and all(len(w) >= 2 for w in words)
        ):
            filtered_names.append(name)

    if not filtered_names:
        possible = re.findall(r'\b[A-Z][a-z]+\b', clean_text)
        filtered_names = possible[:2]

    entities["names"] = list(set(filtered_names))

    # DATES
    dates = re.findall(
        r'\b(?:\d{1,2}\s)?(?:January|February|March|April|May|June|July|'
        r'August|September|October|November|December)\s\d{4}'
        r'|\d{4}\s?[-\u2013]\s?(?:\d{4}|Present)'
        r'|\bGraduated[:\s]+\d{4}\b'
        r'|\b\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}\b',
        clean_text
    )
    entities["dates"] = list(set(dates))

    # ORGANIZATIONS
    orgs = re.findall(
        r'\b[A-Z][A-Za-z\s]*(?:Pvt\.?\s?Ltd\.?|Ltd\.?|Inc\.?|Corp\.?|'
        r'LLP|LLC|University|Company|Agency|Media|School of \w+)\b',
        clean_text
    )
    entities["organizations"] = list(set(orgs))

    # AMOUNTS
    amounts = re.findall(
        r'(?:\u20b9|Rs\.?|\$|\u20ac|\u00a3)\s?\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?'
        r'|\d{1,3}(?:,\d{3})+(?:\.\d{1,2})?',
        clean_text
    )
    entities["amounts"] = list(set(a.strip() for a in amounts if a.strip()))

    # LOCATIONS
    detected = [loc for loc in known_locations if loc in clean_text]
    pins = re.findall(r'\b[1-9][0-9]{5}\b', clean_text)
    city_state = re.findall(r'\b[A-Z][a-z]+,\s[A-Z]{2}\b', clean_text)
    entities["locations"] = list(set(
        detected + city_state + [f"PIN: {p}" for p in pins]
    ))

    # EMAILS
    emails = re.findall(
        r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
        clean_text
    )
    entities["emails"] = list(set(emails))

    # PHONES
    phones = re.findall(
        r'(?:\+?\d{1,3}[\s\-]?)?(?:\(?\d{3}\)?[\s\-]?)?\d{3}[\s\-]?\d{4}',
        clean_text
    )
    entities["phones"] = list(set(phones))

    # WEBSITES
    websites = re.findall(
        r'\b(?:https?://)?(?:www\.)?[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}(?:/\S*)?\b',
        clean_text
    )
    entities["websites"] = list(set(
        [w for w in websites if "@" not in w and "." in w]
    ))

    # SKILLS
    entities["skills"] = list(set(
        [skill for skill in known_skills if skill in clean_text]
    ))

    # DESIGNATIONS
    entities["designations"] = list(set(
        [d for d in known_designations if d in clean_text]
    ))

    return entities

# -------------------------------
# SUMMARY
# -------------------------------
def summarize_text(text):
    if not text:
        return ""
    # Clean extra whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    words = text.split()
    return " ".join(words[:100])

# -------------------------------
# MAIN ANALYSIS
# -------------------------------
def analyze_document(file_name, text):
    entities = extract_entities(text)
    sentiment = get_sentiment(text)
    summary = summarize_text(text)

    # Remove empty fields
    clean_entities = {k: v for k, v in entities.items() if v}

    return {
        "status": "success",
        "fileName": file_name,
        "summary": summary,
        "entities": {
            "names": clean_entities.get("names", []),
            "dates": clean_entities.get("dates", []),
            "organizations": clean_entities.get("organizations", []),
            "amounts": clean_entities.get("amounts", []),
            "locations": clean_entities.get("locations", [])
        },
        "sentiment": sentiment
    }

# -------------------------------
# API ROUTES
# -------------------------------
@app.get("/")
def home():
    return {"message": "API is working 🚀"}

@app.post("/api/document-analyze")
def analyze_api(request: DocumentRequest, x_api_key: str = Header(None)):

    if x_api_key != "mysecret123":
        raise HTTPException(status_code=401, detail="Unauthorized")

    try:
        file_data = base64.b64decode(request.fileBase64)

        # Use temp file instead of saving permanently
        suffix_map = {"pdf": ".pdf", "docx": ".docx", "image": ".png"}
        suffix = suffix_map.get(request.fileType, ".tmp")

        with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
            tmp.write(file_data)
            file_path = tmp.name

        try:
            text = extract_text(file_path, request.fileType)
            result = analyze_document(request.fileName, text)
        finally:
            # Always clean up temp file
            if os.path.exists(file_path):
                os.remove(file_path)

        return result

    except HTTPException:
        raise
    except Exception as e:
        return {"status": "error", "message": str(e)}